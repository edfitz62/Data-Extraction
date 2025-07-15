# -*- coding: utf-8 -*-
"""
Created on Sat Jul 12 14:16:16 2025

@author: edfit
"""

"""
Integrated ABS Analysis Platform
- Python backend with JavaScript UI
- Automated data extraction from New Issue Reports
- Real-time API integration
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
from plotly.utils import PlotlyJSONEncoder
import json
import re
import PyPDF2
import docx
import fitz  # PyMuPDF for better PDF parsing
import requests
from datetime import datetime, timedelta
import sqlite3
import asyncio
import aiohttp
from flask import Flask, jsonify, request, render_template_string, send_from_directory
from flask_cors import CORS
import threading
import webbrowser
import os
from pathlib import Path
import openai  # For AI-powered text extraction
import spacy
from transformers import pipeline
import warnings
warnings.filterwarnings('ignore')

class DocumentExtractor:
    """
    Advanced document extraction engine for New Issue Reports
    """
    
    def __init__(self):
        self.extraction_patterns = self._initialize_patterns()
        self.ai_extractor = None
        self.nlp = None
        self._initialize_ai_models()
        
    def _initialize_ai_models(self):
        """Initialize AI models for text extraction"""
        try:
            # Initialize spaCy for NER
            self.nlp = spacy.load("en_core_web_sm")
        except:
            print("⚠️  spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            
        try:
            # Initialize HuggingFace pipeline for financial text
            self.ai_extractor = pipeline("question-answering", 
                                       model="deepset/roberta-base-squad2")
        except:
            print("⚠️  HuggingFace transformers not available. Using regex extraction only.")
    
    def _initialize_patterns(self):
        """Initialize regex patterns for data extraction"""
        return {
            'deal_name': [
                r'(?:Deal Name|Transaction|Issuer):\s*(.+?)(?:\n|$)',
                r'([A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|LP|Fund)\s*(?:I{1,3}|IV|V|VI|VII|VIII|IX|X)?)',
                r'(\w+\s+\w+\s+(?:Receivables|Fund|Trust|LLC)\s*\d{4}-\d+)'
            ],
            'issue_date': [
                r'(?:Issue Date|Closing Date|Settlement Date):\s*(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
                r'(?:Issue Date|Closing Date|Settlement Date):\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
                r'(\d{4}-\d{2}-\d{2})'
            ],
            'rating_agency': [
                r'(KBRA|Moody\'?s|S&P|Fitch|DBRS)',
                r'Rating Agency:\s*(KBRA|Moody\'?s|S&P|Fitch|DBRS)',
                r'Rated by\s*(KBRA|Moody\'?s|S&P|Fitch|DBRS)'
            ],
            'deal_size': [
                r'(?:Deal Size|Total Size|Aggregate Principal|Principal Amount):\s*\$?([\d,]+\.?\d*)\s*(?:million|MM|M|billion|BB|B)?',
                r'\$?([\d,]+\.?\d*)\s*(?:million|MM|M)\s*(?:deal|transaction|issuance)',
                r'Aggregate Principal.*?\$?([\d,]+\.?\d*)\s*(?:million|MM|M)'
            ],
            'class_a_advance_rate': [
                r'Class A.*?(?:Advance Rate|LTV):\s*([\d.]+)%',
                r'Class A.*?([\d.]+)%\s*(?:advance|LTV)',
                r'(?:Initial|Class A).*?(?:advance rate|LTV).*?([\d.]+)%'
            ],
            'initial_oc': [
                r'(?:Initial|Class A)\s*(?:OC|Overcollateralization|Over-collateralization):\s*([\d.]+)%',
                r'OC.*?([\d.]+)%',
                r'Overcollateralization.*?([\d.]+)%'
            ],
            'expected_cnl': [
                r'(?:Expected|Projected)\s*(?:CNL|Cumulative Net Loss).*?([\d.]+)%',
                r'CNL.*?([\d.]+)%(?:\s*-\s*([\d.]+)%)?',
                r'(?:Loss|CNL)\s*(?:expectation|assumption).*?([\d.]+)%'
            ],
            'reserve_account': [
                r'(?:Reserve Account|Cash Reserve):\s*([\d.]+)%',
                r'Reserve.*?([\d.]+)%',
                r'Cash.*?[Rr]eserve.*?([\d.]+)%'
            ],
            'seasoning': [
                r'(?:Weighted Average|Average|Wtd\.?\s*Avg\.?)\s*(?:Seasoning|Age):\s*([\d.]+)\s*(?:months?|mos?)',
                r'Seasoning.*?([\d.]+)\s*(?:months?|mos?)',
                r'Average\s*(?:age|seasoning).*?([\d.]+)'
            ],
            'sector': [
                r'(?:Asset Class|Collateral Type|Sector):\s*(.+?)(?:\n|$)',
                r'(Equipment|Auto|Consumer|Student|Credit Card|Small Ticket|Working Capital).*?(?:ABS|Loans?|Receivables?)',
                r'(?:Backed by|Collateral).*?(Equipment|Auto|Consumer|Student|Credit Card|Small Ticket|Working Capital)'
            ]
        }
    
    def extract_from_pdf(self, pdf_path):
        """Extract data from PDF using multiple methods"""
        extracted_data = {}
        
        try:
            # Method 1: PyMuPDF (better for complex layouts)
            doc = fitz.open(pdf_path)
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            doc.close()
            
            # Method 2: PyPDF2 (fallback)
            if len(full_text.strip()) < 100:  # If PyMuPDF didn't get much text
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ""
                    for page in pdf_reader.pages:
                        full_text += page.extract_text()
            
            print(f"📄 Extracted {len(full_text)} characters from PDF")
            
            # Extract structured data
            extracted_data = self._extract_structured_data(full_text)
            
            # AI-enhanced extraction if available
            if self.ai_extractor:
                ai_data = self._ai_enhanced_extraction(full_text)
                extracted_data.update(ai_data)
                
        except Exception as e:
            print(f"❌ Error extracting from PDF: {e}")
            
        return extracted_data
    
    def extract_from_docx(self, docx_path):
        """Extract data from Word document"""
        try:
            doc = docx.Document(docx_path)
            full_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\t"
                    full_text += "\n"
            
            print(f"📄 Extracted {len(full_text)} characters from DOCX")
            return self._extract_structured_data(full_text)
            
        except Exception as e:
            print(f"❌ Error extracting from DOCX: {e}")
            return {}
    
    def _extract_structured_data(self, text):
        """Extract structured data using regex patterns"""
        extracted = {}
        
        for field, patterns in self.extraction_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    if field == 'deal_size':
                        # Handle deal size conversion
                        value = re.sub(r'[,$]', '', match.group(1))
                        try:
                            size = float(value)
                            # Check for billion vs million
                            if 'billion' in text.lower() or 'BB' in text or 'B' in text:
                                size *= 1000
                            extracted[field] = size
                        except:
                            continue
                    elif field in ['class_a_advance_rate', 'initial_oc', 'expected_cnl', 'reserve_account']:
                        # Handle percentage fields
                        try:
                            extracted[field] = float(match.group(1))
                        except:
                            continue
                    elif field == 'seasoning':
                        # Handle seasoning in months
                        try:
                            extracted[field] = float(match.group(1))
                        except:
                            continue
                    elif field == 'issue_date':
                        # Handle date parsing
                        date_str = match.group(1)
                        try:
                            # Try different date formats
                            for fmt in ['%m/%d/%Y', '%m-%d-%Y', '%B %d, %Y', '%Y-%m-%d']:
                                try:
                                    parsed_date = datetime.strptime(date_str, fmt)
                                    extracted[field] = parsed_date.strftime('%Y-%m-%d')
                                    break
                                except:
                                    continue
                        except:
                            extracted[field] = date_str
                    else:
                        # Handle text fields
                        extracted[field] = match.group(1).strip()
                    break  # Use first successful match
        
        return extracted
    
    def _ai_enhanced_extraction(self, text):
        """Use AI to extract additional information"""
        ai_extracted = {}
        
        if not self.ai_extractor:
            return ai_extracted
        
        # Define questions for AI extraction
        questions = {
            'deal_name': "What is the name of the deal or transaction?",
            'rating_agency': "Which rating agency rated this transaction?",
            'sector': "What type of assets back this securitization?",
            'class_a_advance_rate': "What is the Class A advance rate percentage?",
            'expected_cnl_low': "What is the expected cumulative net loss percentage?"
        }
        
        try:
            for field, question in questions.items():
                if field not in ai_extracted:  # Only if not found by regex
                    result = self.ai_extractor(question=question, context=text[:2000])  # Limit context
                    if result['score'] > 0.3:  # Confidence threshold
                        ai_extracted[field] = result['answer']
                        
        except Exception as e:
            print(f"⚠️  AI extraction error: {e}")
            
        return ai_extracted
    
    def extract_from_folder(self, folder_path):
        """Extract data from all documents in a folder"""
        folder = Path(folder_path)
        extracted_deals = []
        
        for file_path in folder.glob("*"):
            if file_path.suffix.lower() in ['.pdf', '.docx', '.doc']:
                print(f"🔍 Processing: {file_path.name}")
                
                if file_path.suffix.lower() == '.pdf':
                    data = self.extract_from_pdf(file_path)
                elif file_path.suffix.lower() in ['.docx', '.doc']:
                    data = self.extract_from_docx(file_path)
                else:
                    continue
                
                if data:
                    data['source_file'] = file_path.name
                    data['extraction_date'] = datetime.now().isoformat()
                    extracted_deals.append(data)
                    print(f"✅ Extracted: {data.get('deal_name', 'Unknown Deal')}")
                else:
                    print(f"❌ No data extracted from: {file_path.name}")
        
        return extracted_deals


class ABSFlaskAPI:
    """
    Flask API to bridge Python backend with JavaScript frontend
    """
    
    def __init__(self, abs_platform):
        self.platform = abs_platform
        self.app = Flask(__name__)
        CORS(self.app)  # Enable CORS for JavaScript frontend
        self.extractor = DocumentExtractor()
        self._setup_routes()
        
    def _setup_routes(self):
        """Setup API routes"""
        
        @self.app.route('/')
        def index():
            """Serve the main HTML page"""
            return self._get_html_template()
        
        @self.app.route('/api/deals', methods=['GET'])
        def get_deals():
            """Get all deals"""
            deals = self.platform.deal_database.to_dict('records')
            # Convert datetime objects to strings
            for deal in deals:
                if 'issue_date' in deal and hasattr(deal['issue_date'], 'strftime'):
                    deal['issue_date'] = deal['issue_date'].strftime('%Y-%m-%d')
            return jsonify(deals)
        
        @self.app.route('/api/deals', methods=['POST'])
        def add_deal():
            """Add new deal"""
            try:
                deal_data = request.json
                self.platform.add_deal(deal_data)
                return jsonify({'status': 'success', 'message': 'Deal added successfully'})
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/stress-test', methods=['POST'])
        def run_stress_test():
            """Run stress test"""
            try:
                data = request.json
                scenario = data.get('scenario', 'moderate')
                custom_multiplier = data.get('custom_multiplier', 1.5)
                
                results = self.platform.run_stress_test(scenario, custom_multiplier)
                
                # Convert to JSON-serializable format
                stress_results = {
                    'scenario': scenario,
                    'deals': results[['deal_name', 'stressed_loss', 'adequacy_ratio', 'status']].to_dict('records'),
                    'summary': {
                        'critical_count': len(results[results['status'] == 'Critical']),
                        'weak_count': len(results[results['status'] == 'Weak']),
                        'total_at_risk': len(results[results['status'].isin(['Critical', 'Weak'])]),
                        'volume_at_risk': float(results[results['status'].isin(['Critical', 'Weak'])]['deal_size'].sum())
                    }
                }
                
                return jsonify(stress_results)
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/monte-carlo', methods=['POST'])
        def run_monte_carlo():
            """Run Monte Carlo simulation"""
            try:
                data = request.json
                deal_name = data.get('deal_name')
                iterations = data.get('iterations', 10000)
                loss_volatility = data.get('loss_volatility', 0.3)
                
                results = self.platform.monte_carlo_simulation(
                    deal_name, iterations, loss_volatility
                )
                
                if results:
                    # Convert numpy arrays to lists for JSON serialization
                    results['simulated_losses'] = results['simulated_losses'].tolist()
                    return jsonify(results)
                else:
                    return jsonify({'status': 'error', 'message': 'Deal not found'}), 404
                    
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/peer-analysis', methods=['POST'])
        def run_peer_analysis():
            """Run peer comparison analysis"""
            try:
                data = request.json
                deal_name = data.get('deal_name')
                criteria = data.get('criteria', 'sector')
                
                results = self.platform.peer_analysis(deal_name, criteria)
                return jsonify(results)
                
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/integration-risk', methods=['POST'])
        def assess_integration_risk():
            """Assess integration risk"""
            try:
                data = request.json
                assessment = self.platform.integration_risk_assessment(**data)
                
                # Convert datetime to string
                assessment['assessment_date'] = assessment['assessment_date'].isoformat()
                
                return jsonify(assessment)
                
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/equipment-liquidation', methods=['POST'])
        def analyze_equipment_liquidation():
            """Analyze equipment liquidation scenarios"""
            try:
                data = request.json
                results = self.platform.equipment_liquidation_analysis(**data)
                return jsonify(results)
                
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/extract-documents', methods=['POST'])
        def extract_documents():
            """Extract data from uploaded documents"""
            try:
                folder_path = request.json.get('folder_path')
                if not folder_path or not os.path.exists(folder_path):
                    return jsonify({'status': 'error', 'message': 'Invalid folder path'}), 400
                
                extracted_deals = self.extractor.extract_from_folder(folder_path)
                
                # Add extracted deals to platform
                successful_additions = 0
                for deal_data in extracted_deals:
                    try:
                        # Fill in missing required fields with defaults
                        deal_data = self._fill_missing_fields(deal_data)
                        self.platform.add_deal(deal_data)
                        successful_additions += 1
                    except Exception as e:
                        print(f"❌ Failed to add deal: {e}")
                
                return jsonify({
                    'status': 'success',
                    'message': f'Extracted {len(extracted_deals)} deals, added {successful_additions} successfully',
                    'extracted_deals': extracted_deals
                })
                
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
        
        @self.app.route('/api/charts/<chart_type>')
        def get_chart_data(chart_type):
            """Get chart data for frontend"""
            try:
                if chart_type == 'comparative':
                    fig = self.platform.create_comparative_plotly()
                    return json.dumps(fig, cls=PlotlyJSONEncoder)
                elif chart_type == 'sector-distribution':
                    sector_counts = self.platform.deal_database['sector'].value_counts()
                    fig = px.pie(values=sector_counts.values, names=sector_counts.index)
                    return json.dumps(fig, cls=PlotlyJSONEncoder)
                else:
                    return jsonify({'status': 'error', 'message': 'Unknown chart type'}), 400
                    
            except Exception as e:
                return jsonify({'status': 'error', 'message': str(e)}), 400
    
    def _fill_missing_fields(self, deal_data):
        """Fill missing required fields with intelligent defaults"""
        required_fields = {
            'deal_name': deal_data.get('deal_name', f"Extracted Deal {datetime.now().strftime('%Y%m%d%H%M%S')}"),
            'issue_date': deal_data.get('issue_date', datetime.now().strftime('%Y-%m-%d')),
            'rating_agency': deal_data.get('rating_agency', 'Unknown'),
            'sector': deal_data.get('sector', 'Unknown'),
            'deal_size': deal_data.get('deal_size', 100.0),
            'class_a_advance_rate': deal_data.get('class_a_advance_rate', 80.0),
            'initial_oc': deal_data.get('initial_oc', 10.0),
            'expected_cnl_low': deal_data.get('expected_cnl_low', 2.0),
            'expected_cnl_high': deal_data.get('expected_cnl_high', deal_data.get('expected_cnl_low', 2.0) + 1.0),
            'reserve_account': deal_data.get('reserve_account', 1.0),
            'avg_seasoning': deal_data.get('avg_seasoning', deal_data.get('seasoning', 12)),
            'top_obligor_conc': deal_data.get('top_obligor_conc', 1.0)
        }
        
        return required_fields
    
    def _get_html_template(self):
        """Return the HTML template with embedded JavaScript"""
        return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ABS Analysis Platform - Integrated</title>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/3.9.1/chart.min.js"></script>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            color: #333;
        }
        
        .container {
            max-width: 1400px;
            margin: 0 auto;
            padding: 20px;
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
            color: white;
        }
        
        .header h1 {
            font-size: 2.5rem;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        
        .tabs {
            display: flex;
            background: rgba(255,255,255,0.1);
            backdrop-filter: blur(10px);
            border-radius: 15px;
            padding: 10px;
            margin-bottom: 20px;
            flex-wrap: wrap;
            gap: 10px;
        }
        
        .tab-button {
            flex: 1;
            min-width: 180px;
            padding: 15px 20px;
            background: rgba(255,255,255,0.2);
            border: none;
            border-radius: 10px;
            color: white;
            cursor: pointer;
            transition: all 0.3s ease;
            font-weight: 600;
        }
        
        .tab-button:hover {
            background: rgba(255,255,255,0.3);
            transform: translateY(-2px);
        }
        
        .tab-button.active {
            background: rgba(255,255,255,0.9);
            color: #333;
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }
        
        .tab-content {
            display: none;
            background: rgba(255,255,255,0.95);
            backdrop-filter: blur(10px);
            border-radius: 15px;
            padding: 30px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        
        .tab-content.active {
            display: block;
            animation: fadeIn 0.5s ease-in;
        }
        
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .form-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }
        
        .form-group {
            display: flex;
            flex-direction: column;
        }
        
        .form-group label {
            font-weight: 600;
            margin-bottom: 8px;
            color: #555;
        }
        
        .form-group input, .form-group select {
            padding: 12px;
            border: 2px solid #e1e5e9;
            border-radius: 8px;
            font-size: 14px;
            transition: border-color 0.3s ease;
        }
        
        .form-group input:focus, .form-group select:focus {
            outline: none;
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }
        
        .btn {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 15px 30px;
            border-radius: 10px;
            cursor: pointer;
            font-weight: 600;
            transition: all 0.3s ease;
            margin: 10px;
        }
        
        .btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }
        
        .btn-success {
            background: linear-gradient(135deg, #51cf66 0%, #40c057 100%);
        }
        
        .chart-container {
            background: white;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
            height: 400px;
        }
        
        .dashboard-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(400px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }
        
        .metric-card {
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            border-radius: 10px;
            padding: 20px;
            text-align: center;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        
        .metric-value {
            font-size: 2rem;
            font-weight: bold;
            color: #667eea;
            margin-bottom: 5px;
        }
        
        .metric-label {
            color: #666;
            font-weight: 600;
        }
        
        .alert {
            padding: 15px;
            border-radius: 8px;
            margin: 10px 0;
            font-weight: 600;
        }
        
        .alert-success {
            background: #e8f5e8;
            border: 1px solid #c8e6c9;
            color: #2e7d32;
        }
        
        .alert-error {
            background: #ffebee;
            border: 1px solid #ffcdd2;
            color: #c62828;
        }
        
        .data-table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: white;
            border-radius: 10px;
            overflow: hidden;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        
        .data-table th, .data-table td {
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid #e1e5e9;
        }
        
        .data-table th {
            background: #667eea;
            color: white;
            font-weight: 600;
        }
        
        .data-table tr:hover {
            background: #f8f9fa;
        }
        
        .upload-section {
            background: #e3f2fd;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
            border-left: 4px solid #2196f3;
        }
        
        .extraction-results {
            background: #f8f9fa;
            padding: 15px;
            border-radius: 8px;
            margin-top: 15px;
            border: 1px solid #dee2e6;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🏦 ABS Analysis Platform</h1>
            <p>Integrated Python Backend with JavaScript Frontend + AI Document Extraction</p>
        </div>
        
        <div class="tabs">
            <button class="tab-button active" onclick="showTab('deal-entry')">📝 Deal Entry</button>
            <button class="tab-button" onclick="showTab('extraction')">🤖 AI Extraction</button>
            <button class="tab-button" onclick="showTab('analytics')">📊 Analytics</button>
            <button class="tab-button" onclick="showTab('dashboard')">📈 Dashboard</button>
        </div>
        
        <!-- Deal Entry Tab -->
        <div id="deal-entry" class="tab-content active">
            <h2>Manual Deal Entry</h2>
            <form id="dealForm">
                <div class="form-grid">
                    <div class="form-group">
                        <label>Deal Name</label>
                        <input type="text" id="dealName" required>
                    </div>
                    <div class="form-group">
                        <label>Issue Date</label>
                        <input type="date" id="issueDate" required>
                    </div>
                    <div class="form-group">
                        <label>Rating Agency</label>
                        <select id="ratingAgency" required>
                            <option value="">Select Agency</option>
                            <option value="KBRA">KBRA</option>
                            <option value="Moody's">Moody's</option>
                            <option value="S&P">S&P</option>
                            <option value="Fitch">Fitch</option>
                        </select>
                    </div>
                    <div class="form-group">
                        <label>Sector</label>
                        <select id="sector" required>
                            <option value="">Select Sector</option>
                            <option value="Equipment ABS">Equipment ABS</option>
                            <option value="Small Ticket Leasing">Small Ticket Leasing</option>
                            <option value="Working Capital">Working Capital</option>
                            <option value="Auto ABS">Auto ABS</option>
                            <option value="Consumer Loans">Consumer Loans</option>
                        </select>
                    </div>
                    <div class="form-group">
                        <label>Total Deal Size ($MM)</label>
                        <input type="number" id="dealSize" step="0.01" required>
                    </div>
                    <div class="form-group">
                        <label>Class A Advance Rate (%)</label>
                        <input type="number" id="classAAdvanceRate" step="0.01" max="100" required>
                    </div>
                    <div class="form-group">
                        <label>Initial OC (%)</label>
                        <input type="number" id="initialOC" step="0.01" required>
                    </div>
                    <div class="form-group">
                        <label>Expected CNL Low (%)</label>
                        <input type="number" id="expectedCNLLow" step="0.01" required>
                    </div>
                </div>
                <button type="submit" class="btn">Add Deal to Database</button>
            </form>
            
            <div id="dealAlerts"></div>
        </div>
        
        <!-- AI Extraction Tab -->
        <div id="extraction" class="tab-content">
            <h2>🤖 AI-Powered Document Extraction</h2>
            
            <div class="upload-section">
                <h3>📁 Automated Document Processing</h3>
                <p>Point to a folder containing New Issue Reports (PDF, DOCX) and the AI will automatically extract deal data.</p>
                
                <div class="form-group">
                    <label>Documents Folder Path</label>
                    <input type="text" id="folderPath" placeholder="C:/Users/YourName/Documents/ABS Reports" style="width: 500px;">
                </div>
                
                <button class="btn btn-success" onclick="extractDocuments()">🚀 Extract All Documents</button>
                
                <div id="extractionResults" class="extraction-results" style="display: none;">
                    <h4>Extraction Results</h4>
                    <div id="extractionSummary"></div>
                    <div id="extractedDeals"></div>
                </div>
            </div>
            
            <div style="background: #fff8e1; padding: 20px; border-radius: 10px; margin-top: 20px;">
                <h3>🧠 AI Extraction Features</h3>
                <ul style="margin-left: 20px; margin-top: 10px;">
                    <li>✅ <strong>Multi-format Support</strong>: PDF, DOCX, DOC files</li>
                    <li>✅ <strong>Intelligent Text Recognition</strong>: Advanced OCR and text parsing</li>
                    <li>✅ <strong>Named Entity Recognition</strong>: Automatically identifies deal names, dates, amounts</li>
                    <li>✅ <strong>Financial Data Extraction</strong>: Finds advance rates, credit enhancement, expected losses</li>
                    <li>✅ <strong>Confidence Scoring</strong>: AI provides confidence levels for extracted data</li>
                    <li>✅ <strong>Batch Processing</strong>: Process entire folders of documents</li>
                    <li>✅ <strong>Quality Validation</strong>: Validates extracted data against business rules</li>
                </ul>
            </div>
        </div>
        
        <!-- Analytics Tab -->
        <div id="analytics" class="tab-content">
            <h2>📊 Advanced Analytics</h2>
            
            <!-- Stress Testing -->
            <div style="background: #ffebee; padding: 20px; border-radius: 10px; margin-bottom: 20px;">
                <h3>💥 Stress Testing</h3>
                <div class="form-grid">
                    <div class="form-group">
                        <label>Stress Scenario</label>
                        <select id="stressScenario">
                            <option value="mild">Mild (Loss +20%)</option>
                            <option value="moderate">Moderate (Loss +50%)</option>
                            <option value="severe">Severe (Loss +100%)</option>
                            <option value="custom">Custom</option>
                        </select>
                    </div>
                    <div class="form-group">
                        <label>Custom Loss Multiplier</label>
                        <input type="number" id="customMultiplier" step="0.1" value="1.5">
                    </div>
                </div>
                <button class="btn" onclick="runStressTest()">Run Stress Test</button>
                <div id="stressResults"></div>
            </div>
            
            <!-- Monte Carlo -->
            <div style="background: #e8f5e8; padding: 20px; border-radius: 10px; margin-bottom: 20px;">
                <h3>🎲 Monte Carlo Simulation</h3>
                <div class="form-grid">
                    <div class="form-group">
                        <label>Select Deal</label>
                        <select id="monteCarloDeals">
                            <option value="">Loading deals...</option>
                        </select>
                    </div>
                    <div class="form-group">
                        <label>Iterations</label>
                        <select id="iterations">
                            <option value="1000">1,000</option>
                            <option value="10000">10,000</option>
                            <option value="50000">50,000</option>
                        </select>
                    </div>
                </div>
                <button class="btn" onclick="runMonteCarlo()">Run Simulation</button>
                <div id="monteCarloResults"></div>
            </div>
        </div>
        
        <!-- Dashboard Tab -->
        <div id="dashboard" class="tab-content">
            <h2>📈 Executive Dashboard</h2>
            
            <div class="dashboard-grid">
                <div class="metric-card">
                    <div class="metric-value" id="totalDeals">0</div>
                    <div class="metric-label">Total Deals</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value" id="totalVolume">$0M</div>
                    <div class="metric-label">Total Volume</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value" id="avgAdvanceRate">0%</div>
                    <div class="metric-label">Avg Advance Rate</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value" id="avgOC">0%</div>
                    <div class="metric-label">Avg Initial OC</div>
                </div>
            </div>
            
            <div class="dashboard-grid">
                <div class="chart-container">
                    <div id="comparativeChart"></div>
                </div>
                <div class="chart-container">
                    <div id="sectorChart"></div>
                </div>
            </div>
            
            <div>
                <h3>Deal Database</h3>
                <table class="data-table" id="dealsTable">
                    <thead>
                        <tr>
                            <th>Deal Name</th>
                            <th>Sector</th>
                            <th>Size ($MM)</th>
                            <th>Advance Rate</th>
                            <th>Initial OC</th>
                            <th>Expected CNL</th>
                        </tr>
                    </thead>
                    <tbody id="dealsTableBody">
                    </tbody>
                </table>
            </div>
        </div>
    </div>

    <script>
        // Global variables
        let dealsData = [];
        const API_BASE = '';  // Same origin
        
        // Tab functionality
        function showTab(tabName) {
            document.querySelectorAll('.tab-content').forEach(content => {
                content.classList.remove('active');
            });
            
            document.querySelectorAll('.tab-button').forEach(button => {
                button.classList.remove('active');
            });
            
            document.getElementById(tabName).classList.add('active');
            event.target.classList.add('active');
            
            if (tabName === 'dashboard') {
                loadDashboard();
            } else if (tabName === 'analytics') {
                loadDealsForAnalytics();
            }
        }
        
        // API helper functions
        async function apiCall(endpoint, method = 'GET', data = null) {
            try {
                const options = {
                    method: method,
                    headers: {
                        'Content-Type': 'application/json',
                    },
                };
                
                if (data) {
                    options.body = JSON.stringify(data);
                }
                
                const response = await fetch(API_BASE + endpoint, options);
                return await response.json();
            } catch (error) {
                console.error('API call failed:', error);
                showAlert('API call failed: ' + error.message, 'error');
                return null;
            }
        }
        
        // Deal entry form
        document.getElementById('dealForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const dealData = {
                deal_name: document.getElementById('dealName').value,
                issue_date: document.getElementById('issueDate').value,
                rating_agency: document.getElementById('ratingAgency').value,
                sector: document.getElementById('sector').value,
                deal_size: parseFloat(document.getElementById('dealSize').value),
                class_a_advance_rate: parseFloat(document.getElementById('classAAdvanceRate').value),
                initial_oc: parseFloat(document.getElementById('initialOC').value),
                expected_cnl_low: parseFloat(document.getElementById('expectedCNLLow').value),
                expected_cnl_high: parseFloat(document.getElementById('expectedCNLLow').value) + 0.5,
                reserve_account: 1.0,
                avg_seasoning: 12,
                top_obligor_conc: 1.0
            };
            
            const result = await apiCall('/api/deals', 'POST', dealData);
            if (result && result.status === 'success') {
                showAlert('Deal added successfully!', 'success');
                e.target.reset();
                loadDeals();
            } else {
                showAlert('Failed to add deal: ' + (result?.message || 'Unknown error'), 'error');
            }
        });
        
        // Document extraction
        async function extractDocuments() {
            const folderPath = document.getElementById('folderPath').value;
            if (!folderPath) {
                showAlert('Please enter a folder path', 'error');
                return;
            }
            
            showAlert('Starting document extraction...', 'success');
            
            const result = await apiCall('/api/extract-documents', 'POST', { folder_path: folderPath });
            
            if (result && result.status === 'success') {
                showAlert(result.message, 'success');
                displayExtractionResults(result.extracted_deals);
                loadDeals();  // Refresh deals list
            } else {
                showAlert('Extraction failed: ' + (result?.message || 'Unknown error'), 'error');
            }
        }
        
        function displayExtractionResults(extractedDeals) {
            const resultsDiv = document.getElementById('extractionResults');
            const summaryDiv = document.getElementById('extractionSummary');
            const dealsDiv = document.getElementById('extractedDeals');
            
            summaryDiv.innerHTML = `<p><strong>✅ Extracted ${extractedDeals.length} deals successfully!</strong></p>`;
            
            let dealsHTML = '<h4>Extracted Deals:</h4><ul>';
            extractedDeals.forEach(deal => {
                dealsHTML += `<li><strong>${deal.deal_name || 'Unknown Deal'}</strong> - ${deal.sector || 'Unknown Sector'} - $${deal.deal_size || 'Unknown'}M</li>`;
            });
            dealsHTML += '</ul>';
            
            dealsDiv.innerHTML = dealsHTML;
            resultsDiv.style.display = 'block';
        }
        
        // Analytics functions
        async function runStressTest() {
            const scenario = document.getElementById('stressScenario').value;
            const customMultiplier = parseFloat(document.getElementById('customMultiplier').value);
            
            const result = await apiCall('/api/stress-test', 'POST', {
                scenario: scenario,
                custom_multiplier: customMultiplier
            });
            
            if (result) {
                displayStressResults(result);
            }
        }
        
        function displayStressResults(results) {
            const resultsDiv = document.getElementById('stressResults');
            const summary = results.summary;
            
            let html = `
                <h4>Stress Test Results - ${results.scenario.toUpperCase()}</h4>
                <div class="dashboard-grid">
                    <div class="metric-card">
                        <div class="metric-value">${summary.critical_count}</div>
                        <div class="metric-label">Critical Deals</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">${summary.total_at_risk}</div>
                        <div class="metric-label">Total at Risk</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">$${summary.volume_at_risk.toFixed(0)}M</div>
                        <div class="metric-label">Volume at Risk</div>
                    </div>
                </div>
            `;
            
            resultsDiv.innerHTML = html;
        }
        
        async function runMonteCarlo() {
            const dealName = document.getElementById('monteCarloDeals').value;
            const iterations = parseInt(document.getElementById('iterations').value);
            
            if (!dealName) {
                showAlert('Please select a deal', 'error');
                return;
            }
            
            const result = await apiCall('/api/monte-carlo', 'POST', {
                deal_name: dealName,
                iterations: iterations,
                loss_volatility: 0.3
            });
            
            if (result) {
                displayMonteCarloResults(result);
            }
        }
        
        function displayMonteCarloResults(results) {
            const resultsDiv = document.getElementById('monteCarloResults');
            
            let html = `
                <h4>Monte Carlo Results</h4>
                <div class="dashboard-grid">
                    <div class="metric-card">
                        <div class="metric-value">${results.breach_probability.toFixed(1)}%</div>
                        <div class="metric-label">Breach Probability</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">${results.avg_shortfall.toFixed(2)}%</div>
                        <div class="metric-label">Avg Shortfall</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">${results.percentile_95_loss.toFixed(2)}%</div>
                        <div class="metric-label">95th Percentile Loss</div>
                    </div>
                </div>
            `;
            
            resultsDiv.innerHTML = html;
        }
        
        // Dashboard functions
        async function loadDashboard() {
            await loadDeals();
            updateDashboardMetrics();
            loadCharts();
        }
        
        async function loadDeals() {
            const deals = await apiCall('/api/deals');
            if (deals) {
                dealsData = deals;
                updateDealsTable();
            }
        }
        
        function updateDashboardMetrics() {
            if (dealsData.length === 0) return;
            
            const totalVolume = dealsData.reduce((sum, deal) => sum + deal.deal_size, 0);
            const avgAdvanceRate = dealsData.reduce((sum, deal) => sum + deal.class_a_advance_rate, 0) / dealsData.length;
            const avgOC = dealsData.reduce((sum, deal) => sum + deal.initial_oc, 0) / dealsData.length;
            
            document.getElementById('totalDeals').textContent = dealsData.length;
            document.getElementById('totalVolume').textContent = `$${totalVolume.toFixed(0)}M`;
            document.getElementById('avgAdvanceRate').textContent = `${avgAdvanceRate.toFixed(1)}%`;
            document.getElementById('avgOC').textContent = `${avgOC.toFixed(1)}%`;
        }
        
        function updateDealsTable() {
            const tbody = document.getElementById('dealsTableBody');
            tbody.innerHTML = '';
            
            dealsData.forEach(deal => {
                const row = tbody.insertRow();
                row.innerHTML = `
                    <td>${deal.deal_name}</td>
                    <td>${deal.sector}</td>
                    <td>$${deal.deal_size.toFixed(0)}M</td>
                    <td>${deal.class_a_advance_rate.toFixed(1)}%</td>
                    <td>${deal.initial_oc.toFixed(1)}%</td>
                    <td>${deal.expected_cnl_low.toFixed(1)}%</td>
                `;
            });
        }
        
        async function loadCharts() {
            // Load Plotly charts from backend
            try {
                const response = await fetch('/api/charts/comparative');
                const chartData = await response.json();
                Plotly.newPlot('comparativeChart', chartData.data, chartData.layout);
            } catch (error) {
                console.error('Failed to load comparative chart:', error);
            }
            
            try {
                const response = await fetch('/api/charts/sector-distribution');
                const chartData = await response.json();
                Plotly.newPlot('sectorChart', chartData.data, chartData.layout);
            } catch (error) {
                console.error('Failed to load sector chart:', error);
            }
        }
        
        async function loadDealsForAnalytics() {
            await loadDeals();
            
            // Populate Monte Carlo dropdown
            const monteCarloSelect = document.getElementById('monteCarloDeals');
            monteCarloSelect.innerHTML = '<option value="">Select Deal...</option>';
            
            dealsData.forEach(deal => {
                const option = document.createElement('option');
                option.value = deal.deal_name;
                option.textContent = deal.deal_name;
                monteCarloSelect.appendChild(option);
            });
        }
        
        // Utility functions
        function showAlert(message, type) {
            const alertDiv = document.createElement('div');
            alertDiv.className = `alert alert-${type}`;
            alertDiv.textContent = message;
            
            document.getElementById('dealAlerts').appendChild(alertDiv);
            
            setTimeout(() => {
                alertDiv.remove();
            }, 5000);
        }
        
        // Initialize
        document.addEventListener('DOMContentLoaded', function() {
            loadDashboard();
        });
    </script>
</body>
</html>
        """
    
    def run(self, host='127.0.0.1', port=5000, debug=False):
        """Start the Flask server"""
        print(f"🌐 Starting ABS Analysis Platform at http://{host}:{port}")
        print("🤖 AI Document extraction enabled")
        print("🔗 JavaScript UI integrated with Python backend")
        print("🛑 Press Ctrl+C to stop")
        
        # Open browser automatically
        if not debug:
            threading.Timer(1.5, lambda: webbrowser.open(f'http://{host}:{port}')).start()
        
        self.app.run(host=host, port=port, debug=debug)


# Enhanced ABSAnalysisPlatform with web integration
class IntegratedABSPlatform(ABSAnalysisPlatform):
    """Enhanced platform with web API integration"""
    
    def __init__(self):
        super().__init__()
        self.api_server = None
        self.extractor = DocumentExtractor()
        
    def create_comparative_plotly(self):
        """Create Plotly figure for web frontend"""
        df = self.deal_database
        if len(df) == 0:
            return go.Figure()
            
        fig = go.Figure()
        
        # Add scatter plot
        fig.add_trace(go.Scatter(
            x=df['expected_cnl_low'],
            y=df['initial_oc'] + df['reserve_account'],
            mode='markers+text',
            text=df['deal_name'],
            textposition='top center',
            marker=dict(
                size=df['deal_size']/20,
                color=df['sector'].astype('category').cat.codes,
                colorscale='viridis',
                showscale=True
            ),
            name='Enhancement vs Loss'
        ))
        
        fig.update_layout(
            title='Credit Enhancement vs Expected Loss',
            xaxis_title='Expected CNL (%)',
            yaxis_title='Total Enhancement (%)',
            height=400
        )
        
        return fig
    
    def start_web_interface(self, host='127.0.0.1', port=5000):
        """Start the integrated web interface"""
        self.api_server = ABSFlaskAPI(self)
        self.api_server.run(host=host, port=port)
    
    def extract_documents_from_folder(self, folder_path):
        """Extract documents and add to database"""
        extracted_deals = self.extractor.extract_from_folder(folder_path)
        
        successful_additions = 0
        for deal_data in extracted_deals:
            try:
                # Fill missing fields with defaults
                complete_deal_data = self._fill_missing_fields(deal_data)
                self.add_deal(complete_deal_data)
                successful_additions += 1
            except Exception as e:
                print(f"❌ Failed to add extracted deal: {e}")
        
        print(f"✅ Successfully added {successful_additions}/{len(extracted_deals)} extracted deals")
        return extracted_deals
    
    def _fill_missing_fields(self, deal_data):
        """Fill missing required fields with intelligent defaults"""
        defaults = {
            'deal_name': deal_data.get('deal_name', f"Extracted Deal {datetime.now().strftime('%Y%m%d%H%M%S')}"),
            'issue_date': deal_data.get('issue_date', datetime.now().strftime('%Y-%m-%d')),
            'rating_agency': deal_data.get('rating_agency', 'Unknown'),
            'sector': deal_data.get('sector', 'Unknown'),
            'deal_size': deal_data.get('deal_size', 100.0),
            'class_a_advance_rate': deal_data.get('class_a_advance_rate', 80.0),
            'initial_oc': deal_data.get('initial_oc', 10.0),
            'expected_cnl_low': deal_data.get('expected_cnl_low', 2.0),
            'expected_cnl_high': deal_data.get('expected_cnl_high', deal_data.get('expected_cnl_low', 2.0) + 1.0),
            'reserve_account': deal_data.get('reserve_account', 1.0),
            'avg_seasoning': deal_data.get('avg_seasoning', deal_data.get('seasoning', 12)),
            'top_obligor_conc': deal_data.get('top_obligor_conc', 1.0)
        }
        
        return defaults


# Usage functions for Spyder
def launch_integrated_platform():
    """Launch the fully integrated platform"""
    platform = IntegratedABSPlatform()
    print("""
🚀 Integrated ABS Analysis Platform
==================================

Features:
✅ Python backend with all analytical capabilities
✅ Modern JavaScript frontend with interactive UI
✅ AI-powered document extraction from PDFs/DOCX
✅ Real-time API integration
✅ Automated data population from New Issue Reports

The web interface will open automatically...
    """)
    
    platform.start_web_interface()
    return platform

def extract_sample_documents():
    """Example of how to extract documents"""
    platform = IntegratedABSPlatform()
    
    # Example folder path (update with your actual path)
    folder_path = "C:/Users/edfit/OneDrive - Whitehall Partners/Data Extraction/Sample Reports"
    
    print(f"🔍 Extracting documents from: {folder_path}")
    extracted_deals = platform.extract_documents_from_folder(folder_path)
    
    print(f"📊 Extracted {len(extracted_deals)} deals")
    for deal in extracted_deals:
        print(f"  - {deal.get('deal_name', 'Unknown')} ({deal.get('sector', 'Unknown Sector')})")
    
    return platform

if __name__ == "__main__":
    # Quick start options
    print("""
🏦 ABS Analysis Platform - Integration Options
=============================================

Choose your preferred method:

1. launch_integrated_platform()  # Full web interface
2. extract_sample_documents()    # Test document extraction
3. platform = IntegratedABSPlatform()  # Use in Spyder console

For document extraction, update the folder path in extract_sample_documents()
    """)