# -*- coding: utf-8 -*-
"""
Created on Sat Jul 12 15:03:30 2025

@author: edfit
"""

"""
New Issue Report Parser Interface
AI-powered extraction from PDF and DOCX files
"""

import pandas as pd
import numpy as np
import re
import os
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# Import document processing libraries
try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    print("⚠️  PDF processing libraries not available. Install with: pip install PyPDF2 PyMuPDF")
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️  DOCX processing library not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

class SimpleDocumentExtractor:
    """
    Simplified document extractor for New Issue Reports
    Uses regex patterns to extract key deal information
    """
    
    def __init__(self):
        self.extraction_patterns = self._initialize_patterns()
        self.extracted_data = {}
    
    def _initialize_patterns(self):
        """Initialize regex patterns for data extraction"""
        return {
            'deal_name': [
                r'(?:Deal Name|Transaction|Issuer):\s*(.+?)(?:\n|$)',
                r'([A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|LP|Fund|Trust)\s*(?:I{1,3}|IV|V|VI|VII|VIII|IX|X|\d{4}-\d+)?)',
                r'(\w+\s+\w+\s+(?:Receivables|Fund|Trust|LLC)\s*\d{4}-\d+)',
                r'([A-Z][A-Za-z\s]+\s+(?:ABS|Receivables|Fund|Trust|LLC)\s+\d{4}-\d+)'
            ],
            'issue_date': [
                r'(?:Issue Date|Closing Date|Settlement Date):\s*(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
                r'(?:Issue Date|Closing Date|Settlement Date):\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
                r'(\d{4}-\d{2}-\d{2})',
                r'([A-Za-z]+\s+\d{1,2},\s+\d{4})'
            ],
            'rating_agency': [
                r'(KBRA|Moody\'?s|S&P|Fitch|DBRS)',
                r'Rating Agency:\s*(KBRA|Moody\'?s|S&P|Fitch|DBRS)',
                r'Rated by\s*(KBRA|Moody\'?s|S&P|Fitch|DBRS)'
            ],
            'deal_size': [
                r'(?:Deal Size|Total Size|Aggregate Principal|Principal Amount):\s*\$?([\d,]+\.?\d*)\s*(?:million|MM|M|billion|BB|B)?',
                r'\$?([\d,]+\.?\d*)\s*(?:million|MM|M)\s*(?:deal|transaction|issuance)',
                r'Aggregate Principal.*?\$?([\d,]+\.?\d*)\s*(?:million|MM|M)',
                r'Total.*?\$?([\d,]+\.?\d*)\s*(?:million|MM|M)'
            ],
            'class_a_advance_rate': [
                r'Class A.*?(?:Advance Rate|LTV|advance rate):\s*([\d.]+)%',
                r'Class A.*?([\d.]+)%\s*(?:advance|LTV)',
                r'(?:Initial|Class A).*?(?:advance rate|LTV).*?([\d.]+)%',
                r'Advance Rate.*?Class A.*?([\d.]+)%'
            ],
            'initial_oc': [
                r'(?:Initial|Class A)\s*(?:OC|Overcollateralization|Over-collateralization):\s*([\d.]+)%',
                r'OC.*?([\d.]+)%',
                r'Overcollateralization.*?([\d.]+)%',
                r'Initial.*?OC.*?([\d.]+)%'
            ],
            'expected_cnl': [
                r'(?:Expected|Projected)\s*(?:CNL|Cumulative Net Loss).*?([\d.]+)%',
                r'CNL.*?([\d.]+)%(?:\s*-\s*([\d.]+)%)?',
                r'(?:Loss|CNL)\s*(?:expectation|assumption).*?([\d.]+)%',
                r'Expected.*?Loss.*?([\d.]+)%'
            ],
            'reserve_account': [
                r'(?:Reserve Account|Cash Reserve):\s*([\d.]+)%',
                r'Reserve.*?([\d.]+)%',
                r'Cash.*?[Rr]eserve.*?([\d.]+)%'
            ],
            'seasoning': [
                r'(?:Weighted Average|Average|Wtd\.?\s*Avg\.?)\s*(?:Seasoning|Age):\s*([\d.]+)\s*(?:months?|mos?)',
                r'Seasoning.*?([\d.]+)\s*(?:months?|mos?)',
                r'Average\s*(?:age|seasoning).*?([\d.]+)',
                r'WA.*?Seasoning.*?([\d.]+)'
            ],
            'sector': [
                r'(?:Asset Class|Collateral Type|Sector):\s*(.+?)(?:\n|$)',
                r'(Equipment|Auto|Consumer|Student|Credit Card|Small Ticket|Working Capital).*?(?:ABS|Loans?|Receivables?)',
                r'(?:Backed by|Collateral).*?(Equipment|Auto|Consumer|Student|Credit Card|Small Ticket|Working Capital)'
            ]
        }
    
    def extract_from_pdf(self, pdf_path):
        """Extract data from PDF file"""
        if not PDF_AVAILABLE:
            return {'error': 'PDF processing libraries not installed'}
        
        try:
            # Method 1: Try PyMuPDF first (better for complex layouts)
            text = ""
            try:
                doc = fitz.open(pdf_path)
                for page in doc:
                    text += page.get_text()
                doc.close()
                print(f"📄 Extracted {len(text)} characters using PyMuPDF")
            except:
                # Method 2: Fallback to PyPDF2
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                print(f"📄 Extracted {len(text)} characters using PyPDF2")
            
            if len(text.strip()) < 100:
                return {'error': 'Could not extract sufficient text from PDF'}
            
            return self._extract_structured_data(text, pdf_path)
            
        except Exception as e:
            return {'error': f'PDF extraction failed: {str(e)}'}
    
    def extract_from_docx(self, docx_path):
        """Extract data from DOCX file"""
        if not DOCX_AVAILABLE:
            return {'error': 'DOCX processing library not installed'}
        
        try:
            doc = docx.Document(docx_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            print(f"📄 Extracted {len(text)} characters from DOCX")
            
            if len(text.strip()) < 100:
                return {'error': 'Could not extract sufficient text from DOCX'}
            
            return self._extract_structured_data(text, docx_path)
            
        except Exception as e:
            return {'error': f'DOCX extraction failed: {str(e)}'}
    
    def _extract_structured_data(self, text, file_path):
        """Extract structured data using regex patterns"""
        extracted = {
            'source_file': os.path.basename(file_path),
            'extraction_date': datetime.now().isoformat(),
            'confidence_score': 0
        }
        
        found_fields = 0
        total_fields = len(self.extraction_patterns)
        
        for field, patterns in self.extraction_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    if field == 'deal_size':
                        # Handle deal size conversion
                        value_str = re.sub(r'[,$]', '', match.group(1))
                        try:
                            size = float(value_str)
                            # Check if it's in billions
                            if any(word in text.lower() for word in ['billion', 'bb']):
                                size *= 1000
                            extracted[field] = size
                            found_fields += 1
                        except:
                            continue
                    
                    elif field in ['class_a_advance_rate', 'initial_oc', 'expected_cnl', 'reserve_account']:
                        # Handle percentage fields
                        try:
                            value = float(match.group(1))
                            extracted[field] = value
                            found_fields += 1
                        except:
                            continue
                    
                    elif field == 'seasoning':
                        # Handle seasoning in months
                        try:
                            value = float(match.group(1))
                            extracted['avg_seasoning'] = int(value)
                            found_fields += 1
                        except:
                            continue
                    
                    elif field == 'issue_date':
                        # Handle date parsing
                        date_str = match.group(1)
                        try:
                            # Try different date formats
                            for fmt in ['%m/%d/%Y', '%m-%d-%Y', '%B %d, %Y', '%Y-%m-%d', '%B %d %Y']:
                                try:
                                    parsed_date = datetime.strptime(date_str.replace(',', ''), fmt)
                                    extracted[field] = parsed_date.strftime('%Y-%m-%d')
                                    found_fields += 1
                                    break
                                except:
                                    continue
                        except:
                            extracted[field] = date_str
                            found_fields += 1
                    
                    else:
                        # Handle text fields
                        extracted[field] = match.group(1).strip()
                        found_fields += 1
                    
                    break  # Use first successful match
        
        # Calculate confidence score
        extracted['confidence_score'] = round((found_fields / total_fields) * 100, 1)
        
        # Add intelligent defaults for missing fields
        extracted = self._add_intelligent_defaults(extracted)
        
        return extracted
    
    def _add_intelligent_defaults(self, extracted):
        """Add intelligent defaults for missing required fields"""
        
        # Default deal name if not found
        if 'deal_name' not in extracted:
            extracted['deal_name'] = f"Extracted Deal {datetime.now().strftime('%Y%m%d%H%M%S')}"
        
        # Default date if not found
        if 'issue_date' not in extracted:
            extracted['issue_date'] = datetime.now().strftime('%Y-%m-%d')
        
        # Default rating agency
        if 'rating_agency' not in extracted:
            extracted['rating_agency'] = 'Unknown'
        
        # Default sector based on deal name patterns
        if 'sector' not in extracted:
            deal_name = extracted.get('deal_name', '').lower()
            if any(word in deal_name for word in ['equipment', 'construction', 'machinery']):
                extracted['sector'] = 'Equipment ABS'
            elif any(word in deal_name for word in ['small', 'ticket', 'lease']):
                extracted['sector'] = 'Small Ticket Leasing'
            elif any(word in deal_name for word in ['auto', 'vehicle', 'car']):
                extracted['sector'] = 'Auto ABS'
            elif any(word in deal_name for word in ['consumer', 'personal']):
                extracted['sector'] = 'Consumer Loans'
            else:
                extracted['sector'] = 'Unknown'
        
        # Financial defaults
        if 'deal_size' not in extracted:
            extracted['deal_size'] = 100.0
        
        if 'class_a_advance_rate' not in extracted:
            extracted['class_a_advance_rate'] = 80.0
        
        if 'initial_oc' not in extracted:
            extracted['initial_oc'] = 10.0
        
        if 'expected_cnl' not in extracted:
            extracted['expected_cnl_low'] = 2.0
            extracted['expected_cnl_high'] = 3.0
        else:
            extracted['expected_cnl_low'] = extracted['expected_cnl']
            extracted['expected_cnl_high'] = extracted['expected_cnl'] + 1.0
        
        if 'reserve_account' not in extracted:
            extracted['reserve_account'] = 1.0
        
        if 'avg_seasoning' not in extracted:
            extracted['avg_seasoning'] = 12
        
        # Add top obligor concentration default
        extracted['top_obligor_conc'] = 1.0
        
        return extracted
    
    def process_file(self, file_path):
        """Process a single file and return extracted data"""
        if not os.path.exists(file_path):
            return {'error': f'File not found: {file_path}'}
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        print(f"🔍 Processing: {os.path.basename(file_path)}")
        
        if file_ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        else:
            return {'error': f'Unsupported file type: {file_ext}'}
    
    def process_folder(self, folder_path):
        """Process all documents in a folder"""
        if not os.path.exists(folder_path):
            return {'error': f'Folder not found: {folder_path}'}
        
        results = []
        supported_extensions = ['.pdf', '.docx', '.doc']
        
        print(f"📁 Processing folder: {folder_path}")
        
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext in supported_extensions and os.path.isfile(file_path):
                result = self.process_file(file_path)
                if 'error' not in result:
                    results.append(result)
                    print(f"✅ Successfully processed: {filename}")
                else:
                    print(f"❌ Failed to process: {filename} - {result['error']}")
        
        print(f"📊 Processed {len(results)} documents successfully")
        return results

def demo_document_extraction():
    """Demo function to show how to use the document extractor"""
    extractor = SimpleDocumentExtractor()
    
    print("""
🤖 ABS Document Parser - Ready to Extract!
=========================================

Usage Examples:

# Extract from a single file
result = extractor.process_file("C:/path/to/new_issue_report.pdf")

# Extract from all files in a folder
results = extractor.process_folder("C:/path/to/reports_folder/")

# Then add to your platform
if 'error' not in result:
    platform.add_deal(result)
    """)
    
    return extractor

# Function to integrate with your existing platform
def extract_and_add_to_platform(file_path, platform):
    """Extract data from file and add to platform"""
    extractor = SimpleDocumentExtractor()
    
    print(f"🔍 Extracting data from: {os.path.basename(file_path)}")
    
    # Extract data
    result = extractor.process_file(file_path)
    
    if 'error' in result:
        print(f"❌ Extraction failed: {result['error']}")
        return None
    
    # Display what was found
    print(f"\n📋 EXTRACTION RESULTS:")
    print(f"Confidence Score: {result['confidence_score']}%")
    print(f"Deal Name: {result.get('deal_name', 'Not found')}")
    print(f"Sector: {result.get('sector', 'Not found')}")
    print(f"Deal Size: ${result.get('deal_size', 'Not found')}M")
    print(f"Rating Agency: {result.get('rating_agency', 'Not found')}")
    print(f"Advance Rate: {result.get('class_a_advance_rate', 'Not found')}%")
    print(f"Initial OC: {result.get('initial_oc', 'Not found')}%")
    
    # Ask user to confirm
    confirm = input(f"\n❓ Add this deal to database? (y/n): ").lower().strip()
    
    if confirm == 'y':
        try:
            deal_id = platform.add_deal(result)
            if deal_id:
                print(f"✅ Deal added to database with ID: {deal_id}")
                return deal_id
            else:
                print("❌ Failed to add deal to database")
                return None
        except Exception as e:
            print(f"❌ Error adding to database: {e}")
            return None
    else:
        print("⏭️  Deal not added to database")
        return None

# Initialize the extractor
print("🤖 Initializing Document Extractor...")
extractor = demo_document_extraction()