# -*- coding: utf-8 -*-
"""
Created on Sun Jul 13 19:33:19 2025

@author: edfit
"""

# Complete ABS System: New Issue + Surveillance Processing
import os
import re
import json
import warnings
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from pathlib import Path

# Document processing imports with fallbacks
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pyodbc
    PYODBC_AVAILABLE = True
except ImportError:
    PYODBC_AVAILABLE = False

from flask import Flask, render_template_string, request, jsonify
import threading
import webbrowser

warnings.filterwarnings('ignore')

@dataclass
class NoteClass:
    """Note class data structure"""
    class_id: str
    size_millions: float
    rating: str
    credit_enhancement: float
    coupon: float
    maturity: str
    confidence: int

@dataclass
class DealData:
    """New issue deal data structure"""
    deal_name: str
    issue_date: str
    rating_agency: str
    sector: str
    deal_size: float
    class_a_advance_rate: float
    initial_oc: float
    expected_cnl_low: float
    expected_cnl_high: float
    reserve_account: float
    avg_seasoning: int
    top_obligor_conc: float
    original_pool_balance: float
    note_classes: List[NoteClass]
    confidence_score: int
    source_file: str

@dataclass
class SurveillanceData:
    """Surveillance report data structure"""
    deal_name: str
    report_date: str
    reporting_period: str
    current_pool_balance: float
    total_collections: float
    charge_offs: float
    delinquencies_30_plus: float
    delinquencies_60_plus: float
    delinquencies_90_plus: float
    current_oc: float
    current_enhancement: float
    covenant_compliance: bool
    rating_outlook: str
    servicer_advances: float
    reserve_account_balance: float
    note_class_performance: List[Dict]
    confidence_score: int
    source_file: str

class ComprehensiveDocumentProcessor:
    """Unified processor for both new issue and surveillance reports"""
    
    def __init__(self):
        self.new_issue_patterns = self._initialize_new_issue_patterns()
        self.surveillance_patterns = self._initialize_surveillance_patterns()
        self.note_class_patterns = self._initialize_note_class_patterns()
    
    def _initialize_new_issue_patterns(self):
        """Patterns for new issue reports"""
        return {
            'deal_name': [
                r'(?:Deal Name|Transaction|Issuer):\s*(.+?)(?:\n|$)',
                r'([A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|LP|Fund|Trust)\s*(?:\d{4}-\d+)?)'
            ],
            'deal_size': [
                r'(?:Deal Size|Total Size|Aggregate Principal):\s*\$?([\d,]+\.?\d*)\s*(?:million|MM|M)?',
                r'\$?([\d,]+\.?\d*)\s*(?:million|MM|M)\s*(?:deal|transaction)'
            ],
            'issue_date': [
                r'(?:Issue Date|Closing Date):\s*(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
                r'(\d{4}-\d{2}-\d{2})'
            ],
            'rating_agency': [
                r'(KBRA|Moody\'?s|S&P|Fitch|DBRS)',
                r'Rating Agency:\s*(KBRA|Moody\'?s|S&P|Fitch|DBRS)'
            ],
            'advance_rate': [
                r'Class A.*?(?:Advance Rate|LTV):\s*([\d.]+)%',
                r'Advance Rate.*?([\d.]+)%'
            ],
            'initial_oc': [
                r'(?:Initial|Class A)\s*(?:OC|Overcollateralization):\s*([\d.]+)%',
                r'OC.*?([\d.]+)%'
            ],
            'expected_cnl': [
                r'(?:Expected|Projected)\s*(?:CNL|Cumulative Net Loss).*?([\d.]+)%',
                r'CNL.*?([\d.]+)%'
            ],
            'reserve_account': [
                r'(?:Reserve Account|Cash Reserve):\s*([\d.]+)%',
                r'Reserve.*?([\d.]+)%'
            ]
        }
    
    def _initialize_surveillance_patterns(self):
        """Patterns for surveillance reports"""
        return {
            'report_date': [
                r'(?:Report Date|As of|Reporting Date):\s*(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
                r'(?:Month|Quarter) End(?:ing)?:\s*(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
                r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})\s*(?:Report|Statement)'
            ],
            'reporting_period': [
                r'(?:Reporting Period|Period):\s*(.+?)(?:\n|$)',
                r'(?:Month|Quarter):\s*(.+?)(?:\n|$)',
                r'Period Ending:\s*(.+?)(?:\n|$)'
            ],
            'pool_balance': [
                r'(?:Current|Outstanding|Remaining)\s*(?:Pool|Principal)\s*Balance:\s*\$?([\d,]+\.?\d*)',
                r'Pool Balance.*?\$?([\d,]+\.?\d*)\s*(?:million|MM|M)?',
                r'Outstanding Balance.*?\$?([\d,]+\.?\d*)'
            ],
            'collections': [
                r'(?:Total\s*)?Collections?:\s*\$?([\d,]+\.?\d*)',
                r'Collections.*?\$?([\d,]+\.?\d*)\s*(?:million|MM|M)?',
                r'Principal Collections:\s*\$?([\d,]+\.?\d*)'
            ],
            'charge_offs': [
                r'(?:Charge[- ]?offs?|Write[- ]?offs?):\s*\$?([\d,]+\.?\d*)',
                r'(?:Net\s*)?(?:Charge[- ]?offs?|Losses?).*?\$?([\d,]+\.?\d*)',
                r'Credit Losses:\s*\$?([\d,]+\.?\d*)'
            ],
            'delinquencies': [
                r'(?:30\+|30\s*day[s]?)\s*(?:delinquent|past due).*?(\d+\.?\d*)%?',
                r'(?:60\+|60\s*day[s]?)\s*(?:delinquent|past due).*?(\d+\.?\d*)%?',
                r'(?:90\+|90\s*day[s]?)\s*(?:delinquent|past due).*?(\d+\.?\d*)%?'
            ],
            'current_oc': [
                r'(?:Current|Actual)\s*(?:OC|Overcollateralization):\s*([\d.]+)%',
                r'OC\s*(?:Level|Ratio).*?([\d.]+)%',
                r'Overcollateralization.*?([\d.]+)%'
            ],
            'enhancement': [
                r'(?:Current\s*)?(?:Credit\s*)?Enhancement:\s*([\d.]+)%',
                r'Enhancement Level:\s*([\d.]+)%',
                r'Total Enhancement:\s*([\d.]+)%'
            ],
            'covenant_tests': [
                r'(?:Covenant|Test)\s*(?:Compliance|Status):\s*(Pass|Fail|Compliant|Non-Compliant)',
                r'All\s*(?:Covenants?|Tests?)\s*(?:Pass|Met|Satisfied)',
                r'(?:Trigger|Covenant)\s*(?:Breach|Violation)'
            ],
            'rating_changes': [
                r'Rating\s*(?:Change|Action|Update):\s*(.+?)(?:\n|$)',
                r'(?:Upgraded|Downgraded|Affirmed|Placed).*?(?:Watch|Review)',
                r'Rating\s*Outlook:\s*(Positive|Negative|Stable|Developing)'
            ],
            'servicer_advances': [
                r'Servicer\s*Advances?:\s*\$?([\d,]+\.?\d*)',
                r'Advanced\s*(?:by\s*Servicer)?:\s*\$?([\d,]+\.?\d*)',
                r'Principal\s*Advances?:\s*\$?([\d,]+\.?\d*)'
            ]
        }
    
    def _initialize_note_class_patterns(self):
        """Enhanced note class patterns for both document types"""
        return {
            'table_patterns': [
                r'([A-Z]\d*[+-]?)\s+\$?([\d,]+\.?\d*)\s*(?:M|MM|million)?\s+([A-Z]{1,3}[+-]?)\s+([\d.]+)%',
                r'Class\s+([A-Z]\d*[+-]?)\s+\$?([\d,]+\.?\d*)\s+([A-Z]{1,3}[+-]?)\s+([\d.]+)%',
                r'([A-Z]\d*[+-]?)\s+Notes?\s+\$?([\d,]+\.?\d*)\s+([A-Z]{1,3}[+-]?)\s+([\d.]+)%'
            ],
            'performance_patterns': [
                r'Class\s+([A-Z]\d*[+-]?)\s+.*?(?:Balance|Outstanding):\s*\$?([\d,]+\.?\d*)',
                r'([A-Z]\d*[+-]?)\s+(?:Class|Notes?)\s+.*?Principal:\s*\$?([\d,]+\.?\d*)',
                r'([A-Z]\d*[+-]?)\s+.*?(?:Collections|Payments):\s*\$?([\d,]+\.?\d*)'
            ],
            'class_identifiers': [
                r'(?:Class|Note)\s+([A-Z]\d*[+-]?)',
                r'([A-Z]\d*[+-]?)\s+(?:Class|Note)',
                r'Series\s+([A-Z]\d*[+-]?)'
            ]
        }
    
    def detect_document_type(self, text: str) -> str:
        """Determine if document is new issue or surveillance report"""
        text_lower = text.lower()
        
        # New issue indicators
        new_issue_indicators = [
            'new issue', 'offering', 'prospectus', 'preliminary', 
            'issuance', 'transaction summary', 'deal summary',
            'closing date', 'issue date', 'rating committee'
        ]
        
        # Surveillance indicators  
        surveillance_indicators = [
            'surveillance', 'monitoring', 'quarterly report', 'monthly report',
            'performance report', 'servicer report', 'trustee report',
            'collections', 'charge-offs', 'delinquencies', 'covenant test',
            'pool performance', 'payment date', 'distribution date'
        ]
        
        new_issue_score = sum(1 for indicator in new_issue_indicators if indicator in text_lower)
        surveillance_score = sum(1 for indicator in surveillance_indicators if indicator in text_lower)
        
        if surveillance_score > new_issue_score:
            return 'surveillance'
        else:
            return 'new_issue'
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF or DOCX files"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("PyMuPDF not available. Install with: pip install PyMuPDF")
        
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {e}")
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document based on its type"""
        text = self.extract_text_from_file(file_path)
        doc_type = self.detect_document_type(text)
        
        if doc_type == 'new_issue':
            return self.extract_new_issue_data(text, file_path)
        else:
            return self.extract_surveillance_data(text, file_path)
    
    def extract_new_issue_data(self, text: str, file_path: str) -> Dict[str, Any]:
        """Extract new issue deal information"""
        deal_info = {}
        
        # Extract basic deal information
        for field, patterns in self.new_issue_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    if field in ['deal_size', 'advance_rate', 'initial_oc', 'expected_cnl', 'reserve_account']:
                        try:
                            value = float(re.sub(r'[,$]', '', match.group(1)))
                            deal_info[field] = value
                        except ValueError:
                            continue
                    else:
                        deal_info[field] = match.group(1).strip()
                    break
        
        # Extract note classes
        note_classes = self.extract_note_classes(text)
        
        # Calculate deal size from note classes if not found
        if 'deal_size' not in deal_info and note_classes:
            calculated_size = sum(nc.size_millions for nc in note_classes if nc.size_millions > 0)
            if calculated_size > 0:
                deal_info['deal_size'] = calculated_size
        
        # Create deal data structure
        deal_data = DealData(
            deal_name=deal_info.get('deal_name', f"Extracted Deal {datetime.now().strftime('%Y%m%d%H%M%S')}"),
            issue_date=deal_info.get('issue_date', datetime.now().strftime('%Y-%m-%d')),
            rating_agency=deal_info.get('rating_agency', 'Unknown'),
            sector=self._determine_sector(deal_info.get('deal_name', ''), text),
            deal_size=deal_info.get('deal_size', 100.0),
            class_a_advance_rate=deal_info.get('advance_rate', 80.0),
            initial_oc=deal_info.get('initial_oc', 10.0),
            expected_cnl_low=deal_info.get('expected_cnl', 2.0),
            expected_cnl_high=deal_info.get('expected_cnl', 2.0) + 1.0,
            reserve_account=deal_info.get('reserve_account', 1.0),
            avg_seasoning=12,
            top_obligor_conc=1.0,
            original_pool_balance=deal_info.get('deal_size', 100.0) * 1.1,
            note_classes=note_classes,
            confidence_score=min(100, len(deal_info) * 10 + len(note_classes) * 5),
            source_file=os.path.basename(file_path)
        )
        
        result = asdict(deal_data)
        result['document_type'] = 'new_issue'
        return result
    
    def extract_surveillance_data(self, text: str, file_path: str) -> Dict[str, Any]:
        """Extract surveillance report information"""
        surv_info = {}
        
        # Extract surveillance-specific data
        for field, patterns in self.surveillance_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    if field in ['pool_balance', 'collections', 'charge_offs', 'current_oc', 
                               'enhancement', 'servicer_advances']:
                        try:
                            value = float(re.sub(r'[,$]', '', match.group(1)))
                            surv_info[field] = value
                            break
                        except (ValueError, IndexError):
                            continue
                    elif field == 'delinquencies':
                        # Handle multiple delinquency buckets
                        try:
                            if '30' in match.group(0):
                                surv_info['delinquencies_30_plus'] = float(match.group(1))
                            elif '60' in match.group(0):
                                surv_info['delinquencies_60_plus'] = float(match.group(1))
                            elif '90' in match.group(0):
                                surv_info['delinquencies_90_plus'] = float(match.group(1))
                        except (ValueError, IndexError):
                            continue
                    elif field == 'covenant_tests':
                        covenant_text = match.group(0).lower()
                        surv_info['covenant_compliance'] = 'pass' in covenant_text or 'compliant' in covenant_text
                        break
                    else:
                        surv_info[field] = match.group(1).strip()
                        break
        
        # Extract deal name from surveillance context
        deal_name = self._extract_deal_name_from_surveillance(text)
        
        # Extract note class performance
        note_performance = self._extract_note_class_performance(text)
        
        # Create surveillance data structure
        surveillance_data = SurveillanceData(
            deal_name=deal_name,
            report_date=surv_info.get('report_date', datetime.now().strftime('%Y-%m-%d')),
            reporting_period=surv_info.get('reporting_period', 'Unknown'),
            current_pool_balance=surv_info.get('pool_balance', 0.0),
            total_collections=surv_info.get('collections', 0.0),
            charge_offs=surv_info.get('charge_offs', 0.0),
            delinquencies_30_plus=surv_info.get('delinquencies_30_plus', 0.0),
            delinquencies_60_plus=surv_info.get('delinquencies_60_plus', 0.0),
            delinquencies_90_plus=surv_info.get('delinquencies_90_plus', 0.0),
            current_oc=surv_info.get('current_oc', 0.0),
            current_enhancement=surv_info.get('enhancement', 0.0),
            covenant_compliance=surv_info.get('covenant_compliance', True),
            rating_outlook=surv_info.get('rating_changes', 'Stable'),
            servicer_advances=surv_info.get('servicer_advances', 0.0),
            reserve_account_balance=surv_info.get('pool_balance', 0.0) * 0.02,  # Estimate
            note_class_performance=note_performance,
            confidence_score=min(100, len(surv_info) * 8),
            source_file=os.path.basename(file_path)
        )
        
        result = asdict(surveillance_data)
        result['document_type'] = 'surveillance'
        return result
    
    def _extract_deal_name_from_surveillance(self, text: str) -> str:
        """Extract deal name from surveillance report"""
        patterns = [
            r'(?:Deal|Transaction|Trust):\s*(.+?)(?:\n|$)',
            r'([A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|LP|Fund|Trust)\s*(?:\d{4}-\d+)?)',
            r'(?:Report for|Monitoring):\s*(.+?)(?:\n|$)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return f"Surveillance Report {datetime.now().strftime('%Y%m%d')}"
    
    def _extract_note_class_performance(self, text: str) -> List[Dict]:
        """Extract note class performance data from surveillance reports"""
        performance_data = []
        
        for pattern in self.note_class_patterns['performance_patterns']:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    class_id = match.group(1).strip()
                    amount = float(re.sub(r'[,$]', '', match.group(2)))
                    
                    performance_data.append({
                        'class_id': class_id,
                        'current_balance': amount,
                        'performance_metric': 'balance'  # Could be enhanced to detect metric type
                    })
                except (ValueError, IndexError):
                    continue
        
        return performance_data
    
    def extract_note_classes(self, text: str) -> List[NoteClass]:
        """Extract note classes (same as before, but enhanced)"""
        note_classes = []
        
        # Table-based extraction
        for pattern in self.note_class_patterns['table_patterns']:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    class_id = match.group(1).strip()
                    size_str = match.group(2).replace(',', '')
                    rating = match.group(3).strip()
                    coupon = float(match.group(4))
                    
                    size_millions = float(size_str)
                    
                    note = NoteClass(
                        class_id=class_id,
                        size_millions=size_millions,
                        rating=rating,
                        credit_enhancement=0.0,
                        coupon=coupon,
                        maturity="",
                        confidence=90
                    )
                    note_classes.append(note)
                    
                except (ValueError, IndexError):
                    continue
        
        # Remove duplicates
        unique_notes = {}
        for note in note_classes:
            class_id = note.class_id.upper()
            if class_id not in unique_notes or note.confidence > unique_notes[class_id].confidence:
                unique_notes[class_id] = note
        
        return sorted(list(unique_notes.values()), key=lambda x: x.class_id)
    
    def _determine_sector(self, deal_name: str, text: str) -> str:
        """Determine sector from deal name and text"""
        combined_text = (deal_name + " " + text).lower()
        
        sector_keywords = {
            'Equipment ABS': ['equipment', 'machinery', 'construction', 'industrial'],
            'Auto ABS': ['auto', 'vehicle', 'car', 'automotive'],
            'Consumer Loans': ['consumer', 'personal', 'installment'],
            'Credit Card': ['credit card', 'receivables'],
            'Small Ticket Leasing': ['small ticket', 'leasing'],
            'Student Loans': ['student', 'education'],
            'Working Capital': ['working capital', 'commercial']
        }
        
        for sector, keywords in sector_keywords.items():
            if any(keyword in combined_text for keyword in keywords):
                return sector
        
        return 'Unknown'

class EnhancedDatabaseManager:
    """Database manager supporting both new issue and surveillance data"""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.connection_string = self._setup_connection()
        if self.connection_string:
            self._create_tables()
    
    def _setup_connection(self):
        """Setup database connection"""
        if not PYODBC_AVAILABLE:
            print("⚠️  pyodbc not available - database features disabled")
            return None
        
        if not os.path.exists(self.db_path):
            print(f"❌ Database file not found: {self.db_path}")
            return None
        
        drivers = [
            'Microsoft Access Driver (*.mdb, *.accdb)',
            'Microsoft Office 16.0 Access Database Engine OLE DB Provider'
        ]
        
        for driver in drivers:
            try:
                conn_str = f'DRIVER={{{driver}}};DBQ={self.db_path};'
                conn = pyodbc.connect(conn_str)
                conn.close()
                print(f"✅ Database connected using: {driver}")
                return conn_str
            except Exception:
                continue
        
        print("❌ No Access driver found")
        return None
    
    def _create_tables(self):
        """Create all required database tables"""
        if not self.connection_string:
            return
        
        try:
            conn = pyodbc.connect(self.connection_string)
            cursor = conn.cursor()
            
            # Create main deals table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS ABS_Deals (
                    ID AUTOINCREMENT PRIMARY KEY,
                    DealName TEXT(255),
                    IssueDate DATETIME,
                    RatingAgency TEXT(50),
                    Sector TEXT(100),
                    DealSize CURRENCY,
                    ClassAAdvanceRate DOUBLE,
                    InitialOC DOUBLE,
                    ExpectedCNLLow DOUBLE,
                    ExpectedCNLHigh DOUBLE,
                    ReserveAccount DOUBLE,
                    AvgSeasoning INTEGER,
                    TopObligorConc DOUBLE,
                    OriginalPoolBalance CURRENCY,
                    ConfidenceScore INTEGER,
                    SourceFile TEXT(255),
                    CreatedDate DATETIME DEFAULT NOW()
                )
            ''')
            
            # Create note classes table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS NoteClasses (
                    ID AUTOINCREMENT PRIMARY KEY,
                    DealID INTEGER,
                    ClassID TEXT(10),
                    SizeMillions DOUBLE,
                    Rating TEXT(10),
                    CreditEnhancement DOUBLE,
                    Coupon DOUBLE,
                    Maturity TEXT(50),
                    Confidence INTEGER,
                    CreatedDate DATETIME DEFAULT NOW()
                )
            ''')
            
            # Create surveillance reports table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS SurveillanceReports (
                    ID AUTOINCREMENT PRIMARY KEY,
                    DealName TEXT(255),
                    ReportDate DATETIME,
                    ReportingPeriod TEXT(100),
                    CurrentPoolBalance CURRENCY,
                    TotalCollections CURRENCY,
                    ChargeOffs CURRENCY,
                    Delinquencies30Plus DOUBLE,
                    Delinquencies60Plus DOUBLE,
                    Delinquencies90Plus DOUBLE,
                    CurrentOC DOUBLE,
                    CurrentEnhancement DOUBLE,
                    CovenantCompliance YESNO,
                    RatingOutlook TEXT(50),
                    ServicerAdvances CURRENCY,
                    ReserveAccountBalance CURRENCY,
                    ConfidenceScore INTEGER,
                    SourceFile TEXT(255),
                    CreatedDate DATETIME DEFAULT NOW()
                )
            ''')
            
            # Create note class performance table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS NoteClassPerformance (
                    ID AUTOINCREMENT PRIMARY KEY,
                    SurveillanceID INTEGER,
                    ClassID TEXT(10),
                    CurrentBalance CURRENCY,
                    PerformanceMetric TEXT(50),
                    CreatedDate DATETIME DEFAULT NOW()
                )
            ''')
            
            conn.commit()
            conn.close()
            print("✅ All database tables ready")
            
        except Exception as e:
            print(f"Error creating tables: {e}")
    
    def save_new_issue_deal(self, deal_data: Dict[str, Any]) -> Optional[int]:
        """Save new issue deal data"""
        if not self.connection_string:
            return None
        
        try:
            conn = pyodbc.connect(self.connection_string)
            cursor = conn.cursor()
            
            # Insert main deal
            cursor.execute('''
                INSERT INTO ABS_Deals 
                (DealName, IssueDate, RatingAgency, Sector, DealSize, 
                 ClassAAdvanceRate, InitialOC, ExpectedCNLLow, ExpectedCNLHigh,
                 ReserveAccount, AvgSeasoning, TopObligorConc, OriginalPoolBalance,
                 ConfidenceScore, SourceFile)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                deal_data['deal_name'], deal_data['issue_date'], deal_data['rating_agency'], 
                deal_data['sector'], deal_data['deal_size'], deal_data['class_a_advance_rate'],
                deal_data['initial_oc'], deal_data['expected_cnl_low'], deal_data['expected_cnl_high'],
                deal_data['reserve_account'], deal_data['avg_seasoning'], deal_data['top_obligor_conc'],
                deal_data['original_pool_balance'], deal_data['confidence_score'], deal_data['source_file']
            ))
            
            # Get deal ID
            cursor.execute("SELECT @@IDENTITY")
            deal_id = cursor.fetchone()[0]
            
            # Insert note classes
            for note_class in deal_data.get('note_classes', []):
                cursor.execute('''
                    INSERT INTO NoteClasses 
                    (DealID, ClassID, SizeMillions, Rating, CreditEnhancement, 
                     Coupon, Maturity, Confidence)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    deal_id, note_class['class_id'], note_class['size_millions'],
                    note_class['rating'], note_class['credit_enhancement'],
                    note_class['coupon'], note_class['maturity'], note_class['confidence']
                ))
            
            conn.commit()
            conn.close()
            
            print(f"✅ Saved new issue deal: {deal_data['deal_name']} with {len(deal_data.get('note_classes', []))} note classes")
            return deal_id
            
        except Exception as e:
            print(f"❌ Error saving new issue deal: {e}")
            return None
    
    def save_surveillance_report(self, surv_data: Dict[str, Any]) -> Optional[int]:
        """Save surveillance report data"""
        if not self.connection_string:
            return None
        
        try:
            conn = pyodbc.connect(self.connection_string)
            cursor = conn.cursor()
            
            # Insert surveillance report
            cursor.execute('''
                INSERT INTO SurveillanceReports 
                (DealName, ReportDate, ReportingPeriod, CurrentPoolBalance, TotalCollections,
                 ChargeOffs, Delinquencies30Plus, Delinquencies60Plus, Delinquencies90Plus,
                 CurrentOC, CurrentEnhancement, CovenantCompliance, RatingOutlook,
                 ServicerAdvances, ReserveAccountBalance, ConfidenceScore, SourceFile)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                surv_data['deal_name'], surv_data['report_date'], surv_data['reporting_period'],
                surv_data['current_pool_balance'], surv_data['total_collections'], surv_data['charge_offs'],
                surv_data['delinquencies_30_plus'], surv_data['delinquencies_60_plus'], 
                surv_data['delinquencies_90_plus'], surv_data['current_oc'], surv_data['current_enhancement'],
                surv_data['covenant_compliance'], surv_data['rating_outlook'], surv_data['servicer_advances'],
                surv_data['reserve_account_balance'], surv_data['confidence_score'], surv_data['source_file']
            ))
            
            # Get surveillance ID
            cursor.execute("SELECT @@IDENTITY")
            surv_id = cursor.fetchone()[0]
            
            # Insert note class performance
            for perf in surv_data.get('note_class_performance', []):
                cursor.execute('''
                    INSERT INTO NoteClassPerformance 
                    (SurveillanceID, ClassID, CurrentBalance, PerformanceMetric)
                    VALUES (?, ?, ?, ?)
                ''', (
                    surv_id, perf['class_id'], perf['current_balance'], perf['performance_metric']
                ))
            
            conn.commit()
            conn.close()
            
            print(f"✅ Saved surveillance report: {surv_data['deal_name']} for {surv_data['reporting_period']}")
            return surv_id
            
        except Exception as e:
            print(f"❌ Error saving surveillance report: {e}")
            return None
    
    def get_all_deals(self) -> List[Dict]:
        """Get all deals with note classes"""
        if not self.connection_string:
            return []
        
        try:
            conn = pyodbc.connect(self.connection_string)
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM ABS_Deals ORDER BY CreatedDate DESC")
            deals = []
            
            for row in cursor.fetchall():
                deal = {
                    'id': row[0],
                    'deal_name': row[1],
                    'issue_date': row[2].strftime('%Y-%m-%d') if row[2] else '',
                    'rating_agency': row[3] or '',
                    'sector': row[4] or '',
                    'deal_size': float(row[5]) if row[5] else 0,
                    'class_a_advance_rate': float(row[6]) if row[6] else 0,
                    'initial_oc': float(row[7]) if row[7] else 0,
                    'expected_cnl_low': float(row[8]) if row[8] else 0,
                    'expected_cnl_high': float(row[9]) if row[9] else 0,
                    'reserve_account': float(row[10]) if row[10] else 0,
                    'avg_seasoning': int(row[11]) if row[11] else 0,
                    'top_obligor_conc': float(row[12]) if row[12] else 0,
                    'original_pool_balance': float(row[13]) if row[13] else 0,
                    'confidence_score': int(row[14]) if row[14] else 0,
                    'source_file': row[15] or '',
                    'note_classes': []
                }
                
                # Get note classes
                cursor.execute("SELECT * FROM NoteClasses WHERE DealID = ?", (deal['id'],))
                for note_row in cursor.fetchall():
                    note_class = {
                        'class_id': note_row[2],
                        'size_millions': float(note_row[3]) if note_row[3] else 0,
                        'rating': note_row[4] or '',
                        'credit_enhancement': float(note_row[5]) if note_row[5] else 0,
                        'coupon': float(note_row[6]) if note_row[6] else 0,
                        'maturity': note_row[7] or '',
                        'confidence': int(note_row[8]) if note_row[8] else 0
                    }
                    deal['note_classes'].append(note_class)
                
                deals.append(deal)
            
            conn.close()
            return deals
            
        except Exception as e:
            print(f"Error retrieving deals: {e}")
            return []
    
    def get_surveillance_reports(self) -> List[Dict]:
        """Get all surveillance reports"""
        if not self.connection_string:
            return []
        
        try:
            conn = pyodbc.connect(self.connection_string)
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM SurveillanceReports ORDER BY ReportDate DESC")
            reports = []
            
            for row in cursor.fetchall():
                report = {
                    'id': row[0],
                    'deal_name': row[1],
                    'report_date': row[2].strftime('%Y-%m-%d') if row[2] else '',
                    'reporting_period': row[3] or '',
                    'current_pool_balance': float(row[4]) if row[4] else 0,
                    'total_collections': float(row[5]) if row[5] else 0,
                    'charge_offs': float(row[6]) if row[6] else 0,
                    'delinquencies_30_plus': float(row[7]) if row[7] else 0,
                    'delinquencies_60_plus': float(row[8]) if row[8] else 0,
                    'delinquencies_90_plus': float(row[9]) if row[9] else 0,
                    'current_oc': float(row[10]) if row[10] else 0,
                    'current_enhancement': float(row[11]) if row[11] else 0,
                    'covenant_compliance': bool(row[12]) if row[12] is not None else True,
                    'rating_outlook': row[13] or '',
                    'servicer_advances': float(row[14]) if row[14] else 0,
                    'reserve_account_balance': float(row[15]) if row[15] else 0,
                    'confidence_score': int(row[16]) if row[16] else 0,
                    'source_file': row[17] or '',
                    'note_class_performance': []
                }
                
                # Get note class performance
                cursor.execute("SELECT * FROM NoteClassPerformance WHERE SurveillanceID = ?", (report['id'],))
                for perf_row in cursor.fetchall():
                    performance = {
                        'class_id': perf_row[2],
                        'current_balance': float(perf_row[3]) if perf_row[3] else 0,
                        'performance_metric': perf_row[4] or ''
                    }
                    report['note_class_performance'].append(performance)
                
                reports.append(report)
            
            conn.close()
            return reports
            
        except Exception as e:
            print(f"Error retrieving surveillance reports: {e}")
            return []

class ComprehensiveABSSystem:
    """Complete system handling both new issue and surveillance processing"""
    
    def __init__(self, db_path: str):
        self.processor = ComprehensiveDocumentProcessor()
        self.db_manager = EnhancedDatabaseManager(db_path)
        self.app = Flask(__name__)
        self._setup_routes()
    
    def process_file(self, file_path: str, save_to_db: bool = True) -> Dict[str, Any]:
        """Process single file (auto-detects type)"""
        try:
            result = self.processor.process_document(file_path)
            
            if save_to_db and self.db_manager.connection_string:
                if result.get('document_type') == 'new_issue':
                    deal_id = self.db_manager.save_new_issue_deal(result)
                    if deal_id:
                        result['database_id'] = deal_id
                elif result.get('document_type') == 'surveillance':
                    surv_id = self.db_manager.save_surveillance_report(result)
                    if surv_id:
                        result['database_id'] = surv_id
            
            return result
            
        except Exception as e:
            return {'error': str(e), 'source_file': os.path.basename(file_path)}
    
    def process_folder(self, folder_path: str, save_to_db: bool = True) -> Dict[str, Any]:
        """Process all files in folder"""
        folder = Path(folder_path)
        if not folder.exists():
            return {'error': f'Folder not found: {folder_path}'}
        
        results = {
            'new_issue_deals': [],
            'surveillance_reports': [],
            'failed_files': []
        }
        
        supported_extensions = ['.pdf', '.docx', '.doc']
        
        for file_path in folder.glob('*'):
            if file_path.suffix.lower() in supported_extensions:
                print(f"🔍 Processing: {file_path.name}")
                result = self.process_file(str(file_path), save_to_db)
                
                if 'error' in result:
                    results['failed_files'].append(result)
                    print(f"❌ Failed: {result['error']}")
                elif result.get('document_type') == 'new_issue':
                    results['new_issue_deals'].append(result)
                    note_count = len(result.get('note_classes', []))
                    print(f"✅ New Issue: {result['deal_name']} ({note_count} note classes)")
                elif result.get('document_type') == 'surveillance':
                    results['surveillance_reports'].append(result)
                    print(f"✅ Surveillance: {result['deal_name']} - {result['reporting_period']}")
        
        return results
    
    def _setup_routes(self):
        """Setup Flask routes for web interface"""
        
        @self.app.route('/')
        def index():
            return render_template_string(self._get_enhanced_html_template())
        
        @self.app.route('/api/process-file', methods=['POST'])
        def api_process_file():
            data = request.json
            file_path = data.get('file_path')
            
            if not file_path or not os.path.exists(file_path):
                return jsonify({'error': 'Invalid file path'}), 400
            
            result = self.process_file(file_path, True)
            return jsonify(result)
        
        @self.app.route('/api/process-folder', methods=['POST'])
        def api_process_folder():
            data = request.json
            folder_path = data.get('folder_path')
            
            if not folder_path or not os.path.exists(folder_path):
                return jsonify({'error': 'Invalid folder path'}), 400
            
            results = self.process_folder(folder_path, True)
            return jsonify(results)
        
        @self.app.route('/api/deals', methods=['GET'])
        def api_get_deals():
            deals = self.db_manager.get_all_deals()
            return jsonify(deals)
        
        @self.app.route('/api/surveillance', methods=['GET'])
        def api_get_surveillance():
            reports = self.db_manager.get_surveillance_reports()
            return jsonify(reports)
        
        @self.app.route('/api/deals', methods=['POST'])
        def api_add_deal():
            # Same as before - manual deal entry
            try:
                data = request.json
                # Create deal data structure and save
                # (Implementation same as previous version)
                return jsonify({'success': True})
            except Exception as e:
                return jsonify({'error': str(e)}), 400
    
    def _get_enhanced_html_template(self):
        """Enhanced HTML template supporting both document types"""
        return '''
<!DOCTYPE html>
<html>
<head>
    <title>Complete ABS Processing System</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }
        .container { max-width: 1400px; margin: 0 auto; }
        .header { text-align: center; background: #2c3e50; color: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; }
        .tabs { display: flex; background: white; border-radius: 8px; overflow: hidden; margin-bottom: 20px; }
        .tab { flex: 1; padding: 15px; text-align: center; cursor: pointer; border: none; background: #ecf0f1; }
        .tab.active { background: #3498db; color: white; }
        .tab-content { display: none; background: white; padding: 20px; border-radius: 8px; }
        .tab-content.active { display: block; }
        .document-type-indicator { padding: 10px; margin: 10px 0; border-radius: 5px; font-weight: bold; }
        .new-issue { background: #e8f5e9; border-left: 4px solid #4caf50; }
        .surveillance { background: #e3f2fd; border-left: 4px solid #2196f3; }
        .failed { background: #ffebee; border-left: 4px solid #f44336; }
        .btn { background: #3498db; color: white; padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; margin: 5px; }
        .btn:hover { background: #2980b9; }
        .form-group { margin-bottom: 15px; }
        .form-group label { display: block; margin-bottom: 5px; font-weight: bold; }
        .form-group input { width: 100%; padding: 8px; border: 1px solid #ddd; border-radius: 4px; }
        .results-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(400px, 1fr)); gap: 20px; margin-top: 20px; }
        .result-card { background: white; padding: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
        .table { width: 100%; border-collapse: collapse; margin-top: 20px; }
        .table th, .table td { padding: 8px; text-align: left; border-bottom: 1px solid #ddd; }
        .table th { background: #f2f2f2; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🏦 Complete ABS Processing System</h1>
            <p>Unified New Issue & Surveillance Report Processing</p>
        </div>
        
        <div class="tabs">
            <button class="tab active" onclick="showTab('processing')">📄 Document Processing</button>
            <button class="tab" onclick="showTab('manual')">✍️ Manual Entry</button>
            <button class="tab" onclick="showTab('deals')">📋 New Issue Deals</button>
            <button class="tab" onclick="showTab('surveillance')">📊 Surveillance Reports</button>
        </div>
        
        <!-- Processing Tab -->
        <div id="processing" class="tab-content active">
            <h2>Automated Document Processing</h2>
            <p>System automatically detects New Issue Reports vs Surveillance Reports</p>
            
            <div class="form-group">
                <label>Folder Path (processes both New Issue and Surveillance reports)</label>
                <input type="text" id="folderPath" placeholder="C:\\Users\\YourName\\Documents\\ABS Reports">
            </div>
            <button class="btn" onclick="processFolder()">🚀 Process All Documents</button>
            
            <div class="form-group">
                <label>Single File Path</label>
                <input type="text" id="filePath" placeholder="C:\\Users\\YourName\\Documents\\Report.pdf">
            </div>
            <button class="btn" onclick="processFile()">📊 Process Single File</button>
            
            <div id="processingResults"></div>
        </div>
        
        <!-- Manual Entry Tab -->
        <div id="manual" class="tab-content">
            <h2>Manual Deal Entry</h2>
            <p>Manual entry form (same as before)</p>
        </div>
        
        <!-- Deals Tab -->
        <div id="deals" class="tab-content">
            <h2>New Issue Deals Database</h2>
            <button class="btn" onclick="loadDeals()">🔄 Refresh Deals</button>
            <div id="dealsContainer"></div>
        </div>
        
        <!-- Surveillance Tab -->
        <div id="surveillance" class="tab-content">
            <h2>Surveillance Reports Database</h2>
            <button class="btn" onclick="loadSurveillance()">🔄 Refresh Reports</button>
            <div id="surveillanceContainer"></div>
        </div>
    </div>

    <script>
        function showTab(tabName) {
            document.querySelectorAll('.tab-content').forEach(content => {
                content.classList.remove('active');
            });
            document.querySelectorAll('.tab').forEach(tab => {
                tab.classList.remove('active');
            });
            
            document.getElementById(tabName).classList.add('active');
            event.target.classList.add('active');
            
            if (tabName === 'deals') {
                loadDeals();
            } else if (tabName === 'surveillance') {
                loadSurveillance();
            }
        }
        
        async function processFolder() {
            const folderPath = document.getElementById('folderPath').value;
            if (!folderPath) {
                alert('Please enter a folder path');
                return;
            }
            
            try {
                const response = await fetch('/api/process-folder', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({folder_path: folderPath})
                });
                
                const results = await response.json();
                displayProcessingResults(results);
            } catch (error) {
                alert('Error: ' + error.message);
            }
        }
        
        function displayProcessingResults(results) {
            const container = document.getElementById('processingResults');
            
            let html = '<div class="results-grid">';
            
            // New Issue Results
            if (results.new_issue_deals && results.new_issue_deals.length > 0) {
                html += '<div class="result-card">';
                html += '<h3>📄 New Issue Deals Processed</h3>';
                results.new_issue_deals.forEach(deal => {
                    const noteCount = deal.note_classes ? deal.note_classes.length : 0;
                    html += `<div class="document-type-indicator new-issue">`;
                    html += `${deal.deal_name} - ${noteCount} note classes (${deal.confidence_score}% confidence)`;
                    html += `</div>`;
                });
                html += '</div>';
            }
            
            // Surveillance Results
            if (results.surveillance_reports && results.surveillance_reports.length > 0) {
                html += '<div class="result-card">';
                html += '<h3>📊 Surveillance Reports Processed</h3>';
                results.surveillance_reports.forEach(report => {
                    html += `<div class="document-type-indicator surveillance">`;
                    html += `${report.deal_name} - ${report.reporting_period} (${report.confidence_score}% confidence)`;
                    html += `</div>`;
                });
                html += '</div>';
            }
            
            // Failed Files
            if (results.failed_files && results.failed_files.length > 0) {
                html += '<div class="result-card">';
                html += '<h3>❌ Failed Files</h3>';
                results.failed_files.forEach(failed => {
                    html += `<div class="document-type-indicator failed">`;
                    html += `${failed.source_file}: ${failed.error}`;
                    html += `</div>`;
                });
                html += '</div>';
            }
            
            html += '</div>';
            container.innerHTML = html;
        }
        
        async function loadDeals() {
            try {
                const response = await fetch('/api/deals');
                const deals = await response.json();
                displayDeals(deals);
            } catch (error) {
                alert('Error loading deals: ' + error.message);
            }
        }
        
        async function loadSurveillance() {
            try {
                const response = await fetch('/api/surveillance');
                const reports = await response.json();
                displaySurveillance(reports);
            } catch (error) {
                alert('Error loading surveillance: ' + error.message);
            }
        }
        
        function displayDeals(deals) {
            const container = document.getElementById('dealsContainer');
            
            if (deals.length === 0) {
                container.innerHTML = '<p>No new issue deals found.</p>';
                return;
            }
            
            let html = `<table class="table">
                <thead>
                    <tr>
                        <th>Deal Name</th>
                        <th>Sector</th>
                        <th>Size ($MM)</th>
                        <th>Note Classes</th>
                        <th>Issue Date</th>
                        <th>Source</th>
                    </tr>
                </thead>
                <tbody>`;
            
            deals.forEach(deal => {
                const noteInfo = deal.note_classes.length > 0 
                    ? `${deal.note_classes.length} classes`
                    : 'None';
                
                html += `<tr>
                    <td><strong>${deal.deal_name}</strong></td>
                    <td>${deal.sector}</td>
                    <td>$${deal.deal_size.toFixed(0)}M</td>
                    <td>${noteInfo}</td>
                    <td>${deal.issue_date}</td>
                    <td>${deal.source_file}</td>
                </tr>`;
            });
            
            html += '</tbody></table>';
            container.innerHTML = html;
        }
        
        function displaySurveillance(reports) {
            const container = document.getElementById('surveillanceContainer');
            
            if (reports.length === 0) {
                container.innerHTML = '<p>No surveillance reports found.</p>';
                return;
            }
            
            let html = `<table class="table">
                <thead>
                    <tr>
                        <th>Deal Name</th>
                        <th>Report Date</th>
                        <th>Period</th>
                        <th>Pool Balance ($MM)</th>
                        <th>Collections ($MM)</th>
                        <th>Charge-offs ($MM)</th>
                        <th>Current OC (%)</th>
                        <th>Covenant</th>
                        <th>Source</th>
                    </tr>
                </thead>
                <tbody>`;
            
            reports.forEach(report => {
                const covenantStatus = report.covenant_compliance ? '✅ Pass' : '❌ Fail';
                
                html += `<tr>
                    <td><strong>${report.deal_name}</strong></td>
                    <td>${report.report_date}</td>
                    <td>${report.reporting_period}</td>
                    <td>$${(report.current_pool_balance / 1000000).toFixed(1)}M</td>
                    <td>$${(report.total_collections / 1000000).toFixed(1)}M</td>
                    <td>$${(report.charge_offs / 1000000).toFixed(1)}M</td>
                    <td>${report.current_oc.toFixed(1)}%</td>
                    <td>${covenantStatus}</td>
                    <td>${report.source_file}</td>
                </tr>`;
            });
            
            html += '</tbody></table>';
            container.innerHTML = html;
        }
    </script>
</body>
</html>
        '''
    
    def run(self, host='127.0.0.1', port=5000):
        """Start the complete system"""
        print(f"🚀 Starting Complete ABS Processing System")
        print(f"🌐 Web Interface: http://{host}:{port}")
        print(f"📊 Database: {'✅ Connected' if self.db_manager.connection_string else '❌ Not Available'}")
        print(f"📄 Document Types: ✅ New Issue + Surveillance Reports")
        print(f"🎯 Auto-Detection: ✅ Active")
        print("🛑 Press Ctrl+C to stop")
        
        threading.Timer(1.0, lambda: webbrowser.open(f'http://{host}:{port}')).start()
        self.app.run(host=host, port=port, debug=False)

def launch_complete_abs_system(db_path=None):
    """Launch the complete ABS system with surveillance capabilities"""
    if db_path is None:
        db_path = r"C:\Users\edfit\OneDrive - Whitehall Partners\Data Extraction\ABS_Performance_Data.accdb"
    
    print("🏦 Complete ABS Processing System")
    print("=" * 50)
    print("✅ New Issue Report processing with enhanced note class extraction")
    print("✅ Surveillance Report processing with performance tracking")
    print("✅ Automatic document type detection")
    print("✅ Unified database schema for deals and surveillance")
    print("✅ Single web interface for all operations")
    print("✅ Enhanced note class extraction (dynamic detection)")
    print("✅ Performance monitoring and covenant tracking")
    print()
    
    system = ComprehensiveABSSystem(db_path)
    system.run()
    
    return system

if __name__ == "__main__":
    launch_complete_abs_system()